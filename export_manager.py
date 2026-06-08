"""
Export Manager
Handles PDF and Word document creation with enhanced error handling.
"""

import logging
from io import BytesIO
from typing import Dict

logger = logging.getLogger(__name__)

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class ExportManager:
    """Handles PDF and Word document exports."""
    
    @staticmethod
    def create_pdf(content: str, title: str, metadata: Dict) -> BytesIO:
        """Create PDF document with proper error handling."""
        buffer = BytesIO()
        
        if not REPORTLAB_AVAILABLE:
            return ExportManager._create_text_fallback(content, title, metadata)
        
        try:
            doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=72, bottomMargin=72)
            
            styles = getSampleStyleSheet()
            story = []
            
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=TA_CENTER
            )
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 12))
            
            for key, value in metadata.items():
                story.append(Paragraph(f"<b>{key}:</b> {value}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            content_lines = content.split('\n')
            for line in content_lines:
                if line.strip():
                    if line.startswith('#'):
                        level = len(line.split('#')) - 1
                        header_text = line.replace('#', '').strip()
                        if level == 1:
                            story.append(Paragraph(header_text, styles['Heading1']))
                        elif level == 2:
                            story.append(Paragraph(header_text, styles['Heading2']))
                        else:
                            story.append(Paragraph(header_text, styles['Heading3']))
                    else:
                        escaped_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        story.append(Paragraph(escaped_line, styles['Normal']))
                story.append(Spacer(1, 6))
            
            doc.build(story)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            logger.error(f"PDF creation error: {e}")
            return ExportManager._create_text_fallback(content, title, metadata)
    
    @staticmethod
    def create_word_doc(content: str, title: str, metadata: Dict) -> BytesIO:
        """Create Word document with proper error handling."""
        buffer = BytesIO()
        
        if not DOCX_AVAILABLE:
            return ExportManager._create_text_fallback(content, title, metadata)
        
        try:
            doc = Document()
            
            title_para = doc.add_heading(title, 0)
            title_para.alignment = 1
            
            doc.add_paragraph()
            for key, value in metadata.items():
                p = doc.add_paragraph()
                p.add_run(f"{key}: ").bold = True
                p.add_run(str(value))
            
            doc.add_paragraph()
            
            content_lines = content.split('\n')
            for line in content_lines:
                if line.strip():
                    if line.startswith('#'):
                        level = len(line.split('#')) - 1
                        header_text = line.replace('#', '').strip()
                        doc.add_heading(header_text, min(level, 6))
                    else:
                        doc.add_paragraph(line)
            
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except Exception as e:
            logger.error(f"Word document creation error: {e}")
            return ExportManager._create_text_fallback(content, title, metadata)
    
    @staticmethod
    def _create_text_fallback(content: str, title: str, metadata: Dict) -> BytesIO:
        """Create text file fallback."""
        buffer = BytesIO()
        
        text_content = f"{title}\n" + "="*len(title) + "\n\n"
        
        for key, value in metadata.items():
            text_content += f"{key}: {value}\n"
        
        text_content += f"\n{content}"
        
        buffer.write(text_content.encode('utf-8'))
        buffer.seek(0)
        return buffer
    
    @staticmethod
    def get_file_extension(doc_type: str) -> str:
        """Get appropriate file extension."""
        if doc_type == 'pdf':
            return "pdf" if REPORTLAB_AVAILABLE else "txt"
        elif doc_type == 'word':
            return "docx" if DOCX_AVAILABLE else "txt"
        else:
            return "txt"
    
    @staticmethod
    def get_mime_type(doc_type: str) -> str:
        """Get appropriate MIME type."""
        if doc_type == 'pdf':
            return "application/pdf" if REPORTLAB_AVAILABLE else "text/plain"
        elif doc_type == 'word':
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if DOCX_AVAILABLE else "text/plain"
        else:
            return "text/plain"
